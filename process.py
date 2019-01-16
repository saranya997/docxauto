
from docx import Document
from docx.shared import Inches
import re
import xml.etree.ElementTree as ET
from PIL import Image
import os
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from nltk.corpus import stopwords 
from nltk.tokenize import word_tokenize 
import string

#filters punctuation and stopwords
def cleanSearch(text):
    tokens = word_tokenize(text)
    tokens = [w.lower() for w in tokens]
    table = str.maketrans('', '', string.punctuation)
    stripped = [w.translate(table) for w in tokens]
    words = [word for word in stripped if word.isalpha()]
    stop_words = set(stopwords.words('english'))
    words = [w for w in words if not w in stop_words]
    return words



def searchdoc(search_query):
    filename="load.docx"
    doc_in=Document('./uploads/'+filename)
    doc_out=Document()

    # section_list is used to store paragraph,images,tabloes which are comes under heading which contains search words
    sections_list = []
    #bool to maintain search status
    search_flag = 0
    #para_section_list is used to store only paragraphs
    para_section_list=[]

    # section_heading is used to store headings
    section_heading = None

    #section_paragraphs is used to store paragraphs
    section_paragraphs = []

    #section_table is used to store table
    section_table=[]

    #section_pic is used to store images
    section_pic=[]

    #para_flag which indicates paragraph is heading
    para_flag=0

    # table_flag which indicates table is there
    table_flag=0

    #i which indicates one section is completed 
    i=0

    #next_paragraph_style_number which tells style size of next paragraph 
    next_paragraph_style_number=10

    #table_heading_only  which indicates table heading 
    table_heading_only=0

    search_content=cleanSearch(search_query)
    print(search_content)
    #for image identification

    def hasImage(par):
        ids = []
        root = ET.fromstring(par._p.xml)
        namespace = {
                 'a':"http://schemas.openxmlformats.org/drawingml/2006/main", \
                 'r':"http://schemas.openxmlformats.org/officeDocument/2006/relationships", \
                 'wp':"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}

        inlines = root.findall('.//wp:inline',namespace)
        for inline in inlines:
            imgs = inline.findall('.//a:blip', namespace)
            for img in imgs:     
                id = img.attrib['{{{0}}}embed'.format(namespace['r'])]
            ids.append(id)
        return ids


    # for table content and paragarph content identification



    def iter_block_items(parent):
     if isinstance(parent, _Document):
        parent_elm = parent.element.body
     elif isinstance(parent, _Cell):
        parent_elm = parent._tc
     elif isinstance(parent, _Row):
        parent_elm = parent._tr
     for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

    # table and paragraph content extraction


    for paragraph in iter_block_items(doc_in):
     
    # table content extraction 

     if isinstance(paragraph,Table) and table_flag ==1:
      for row in paragraph.rows:
            rowdata=" "
            for cell in row.cells:
                
                for paragraph in cell.paragraphs:
                    rowdata=rowdata+"  "+pargraph.text
            section_table.append(rowdata)
      section_table.append("TOCend")         
     
    # paragraph content extraction

     if isinstance(paragraph, Paragraph):
      
    #for title 
      
      if any(ext in paragraph.text.lower() for ext in search_content) and paragraph.style.name.startswith('Title'):
        search_flag = 1
	
        doc_out.add_paragraph("word is in title so entire document deals with that word",style="heading 1")
      
    # for heading and sub heading
      
      if any(ext in paragraph.text.lower() for ext in search_content) or para_flag == 1:
       search_flag = 1
       if paragraph.style.name.startswith('Heading') or para_flag == 1:
        if paragraph.style.name.startswith('Heading'):
         next_paragraph_style_name=paragraph.style.name
         next_paragraph_style_number=re.findall('\d+',next_paragraph_style_name)
        
        if para_flag == 0 or any(ext in paragraph.text.lower() for ext in search_content):
         current_paragraph_style_name=paragraph.style.name
         current_paragraph_style_number=re.findall('\d+',current_paragraph_style_name)
        
        if paragraph.style.name.startswith(current_paragraph_style_name) or current_paragraph_style_number > next_paragraph_style_number :
        
         section = {
                'heading': section_heading,
                'paragraphs': section_paragraphs,
                'table': section_table,
                'picture':   section_pic
         }
         sections_list.append(section)
        
         i=i+1
         section_heading = paragraph.text
         table_flag=1
         para_flag=1
         if i>1:
          i=0
          para_flag=0
          table_flag=0
         section_paragraphs = []
         section_table=[]
         section_pic=[]
         continue
        if current_paragraph_style_number <= next_paragraph_style_number:
          section_paragraphs.append(paragraph)
          for id in hasImage(paragraph): 
           rID = id
           document_part = doc_in.part
           image_part = document_part.related_parts[rID]
           section_pic.append(image_part)
       else:
        if paragraph.style.name != "Title":
         para_section_list.append(paragraph.text)


    #copy heading content in docx



    for section in sections_list:   
        doc_out.add_heading(section['heading'])
        for paragraph in section['paragraphs']:
         if paragraph.style.name.startswith('Heading'):
           paragraph_style=paragraph.style.name
           doc_out.add_paragraph(paragraph.text,style=paragraph_style)
           
           
         else:
          doc_out.add_paragraph(paragraph.text)
        for paragraph in section['table']:
          if table_heading_only==0:
           doc_out.add_paragraph(paragraph,style='heading 5')
           table_heading_only=table_heading_only+1
          else:
           if "TOCend" in paragraph:
            table_heading_only=0
            continue
           else:
            doc_out.add_paragraph(paragraph)
        
        

    # copy paragraph content only



    doc_out.add_paragraph(para_section_list) 


    # save docx    
    doc_out.save('output/output.docx')
    if search_flag:
        return "found"
    else:
        return "not found"

    #delete image in user desktop   

    #os.remove("im.png")    


